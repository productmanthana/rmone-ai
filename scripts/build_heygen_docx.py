#!/usr/bin/env python3
"""Build the HeyGen walkthrough script as a Word doc with embedded screenshots."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "screenshots"
OUT = ROOT / "RM_ONE_HeyGen_Script.docx"

# Scene = (number, title, runtime, shot_file, voiceover, on_screen_text)
SCENES = [
    (1, "Cold Open — The Command Centre", "0:00 – 0:14", "01-home.jpg",
     "This is RM ONE. Not a project tracker. Not a scheduling tool. "
     "An operational command centre — the single screen where every signal "
     "across LiRo's business arrives, gets ranked, and turns into a decision. "
     "Live. Right now.",
     "RM ONE · OPERATIONAL COMMAND CENTRE"),

    (2, "Mission Control (Home)", "0:14 – 0:38", "01-home.jpg",
     "When you sign in, RM ONE shows you the state of the business — not a "
     "to-do list. Seven-point-eight million in active value. Fifty-seven "
     "million in pipeline. Thirty-five qualified leads waiting on a decision. "
     "Below the headline, the top roles in demand and the largest active "
     "engagements ranked by exposure. This isn't a report you run. It's the "
     "operational pulse — refreshed continuously from your live systems.",
     "EXECUTIVE SNAPSHOT · LIVE PULSE"),

    (3, "Engagement Pipeline", "0:38 – 1:00", "02-projects.jpg",
     "From mission control, drill into the engagement pipeline. Two-hundred-"
     "seventy-one active engagements across the firm, every one of them a "
     "live card showing phase, contract value, schedule drift, and the "
     "actions you can take right now — view, reallocate with AI, or jump "
     "to the team. No spreadsheets. No status meetings. Just operational "
     "state.",
     "271 ACTIVE · LIVE STATUS · ZERO LATENCY"),

    (4, "Opportunities · Leads · Accounts", "1:00 – 1:28", "12-projects-opps.jpg",
     "The same command centre, three more lenses. Opportunities is your "
     "forward funnel. Leads — seventeen-hundred-sixty-five of them — surface "
     "the largest unsigned exposures first; an eighteen-million-dollar "
     "Townsend deal, a two-hundred-million-dollar Bio-Pharm engagement, "
     "Google's hundred-fifty-million campus. Accounts gives you thirty-"
     "eight-hundred companies, three thousand active customers, each with "
     "their full engagement history and an AI-generated profile one click "
     "away.",
     "PIPELINE · LEADS · ACCOUNTS"),

    (5, "Leads at a Glance", "1:28 – 1:40", "13-projects-leads.jpg",
     "Every lead ranked by value, every action one tap away — Pursue, "
     "Pre-Staff, Details. The command centre doesn't ask you to remember "
     "what's hot. It tells you.",
     "1,765 LEADS · RANKED BY VALUE"),

    (6, "Account Intelligence", "1:40 – 1:52", "14-projects-companies.jpg",
     "Three-thousand-eight-hundred-thirty-one companies on file. Two-"
     "thousand-nine-hundred-ninety-eight active customers. Every account "
     "carries its full operational history — engagements, contacts, "
     "AI-generated relationship profile.",
     "3,831 ACCOUNTS · 2,998 ACTIVE"),

    (7, "Engagement Deep-Dive", "1:52 – 2:10", "11-project-detail.jpg",
     "Open any engagement and the command centre orchestrates the lookup — "
     "the record, the team allocations, the financials, the health gauge, "
     "the dashboard — pulled and reconciled the moment you open it. "
     "Nothing is stale. Nothing is exported. Everything is live.",
     "LIVE RECONCILIATION · NO STALE DATA"),

    (8, "Workforce Capacity — Timeline", "2:10 – 2:34", "03-resources.jpg",
     "The heart of the operational command centre is the workforce view. "
     "Fifty-three operators. Thirteen weeks of forward capacity. Colour-"
     "coded — green is optimal, amber is under-loaded, red is overloaded. "
     "Right now, Bruce Korrow sits at one-hundred-thirty percent. Joseph "
     "Massa is fully available at zero. The command centre tells you who's "
     "drowning and who's free, instantly.",
     "53 OPERATORS · 13-WEEK FORWARD VIEW"),

    (9, "Staff View — Allocation Gauges", "2:34 – 2:48", "15-resources-staff.jpg",
     "Switch lenses. Staff view shows every operator as a card with an "
     "allocation gauge — overloaded, optimal, active, under-used — and "
     "one-click access to view, assign, or open their profile. The "
     "command centre makes reallocation a single decision, not a meeting.",
     "ALLOCATION GAUGES · ONE-CLICK ASSIGN"),

    (10, "Demand Signals · AI Staffing", "2:48 – 3:04", "16-resources-demand.jpg",
     "Demand surfaces every open role — one-hundred-three of them right "
     "now — with a Find Staff with AI button on each. Click it and the "
     "command centre scores your entire bench against the role in real "
     "time. The right operator, surfaced automatically.",
     "103 OPEN DEMAND · AI MATCHING"),

    (11, "Relationships — CRM", "3:04 – 3:14", "17-resources-crm.jpg",
     "Client and partner contacts live in the same command centre. No "
     "external CRM. No data silo. Same screen, same source of truth.",
     "UNIFIED CONTACTS · NO SILOS"),

    (12, "AI Decision Support", "3:14 – 3:40", "04-chat.jpg",
     "This is where the command centre becomes an operator. Ask anything "
     "of your live data — who is under-utilised, show me bench resources, "
     "pipeline health summary — and the AI answers from the operational "
     "ground truth. It can also draft and send email on your behalf, right "
     "from the conversation. Voice dictation built in. Ninety-nine plus "
     "live signals tracked in the background, every minute of every day.",
     "AI · VOICE · EMAIL · 99+ LIVE SIGNALS"),

    (13, "Operator Inbox", "3:40 – 3:54", "18-chat-inbox.jpg",
     "Every email the command centre has sent on your behalf — and every "
     "reply waiting on you — surfaces in the operator inbox. Received, "
     "sent, all. One queue. Zero context-switching.",
     "OPERATOR INBOX · RECEIVED · SENT"),

    (14, "Risk Feed (Alerts)", "3:54 – 4:14", "05-alerts.jpg",
     "The command centre's risk feed runs continuously. Eleven live "
     "signals at this moment — three critical, eight warning. Bruce Korrow "
     "over-allocated at one-thirty across forty-two engagements. Twenty "
     "operators on the bench. Twenty-three open staffing demands. Bench "
     "headcount climbing week-over-week. Every operational risk, on one "
     "screen, ranked by severity.",
     "11 LIVE SIGNALS · RANKED BY SEVERITY"),

    (15, "Operational Analytics", "4:14 – 4:34", "06-analytics.jpg",
     "Analytics rolls the operating picture up. Eight million in open "
     "engagement value. Seventy-one million in long-term leads. Client "
     "concentration tells you where you're exposed — NYC Parks at thirty-"
     "one percent, NYC Housing at twenty-four, Con Edison at twenty-one. "
     "This is your weekly leadership read in a single scroll.",
     "EXPOSURE · CONCENTRATION · MIX"),

    (16, "Cost Source-of-Truth (Rate Card)", "4:34 – 4:44", "08-rate-card.jpg",
     "Every cost decision in the command centre is anchored to the rate "
     "card — per office, per role, per job title. Change a rate here and "
     "it flows into every new engagement budget automatically.",
     "RATE CARD · SOURCE OF TRUTH"),

    (17, "Self-Monitoring (System Health)", "4:44 – 4:58", "10-system-health.jpg",
     "Because this is operational infrastructure, the command centre "
     "monitors itself. One-hundred percent uptime. Sixty-millisecond "
     "average latency. Three of three services operational. You always "
     "know the platform is healthy — because if it isn't, neither is the "
     "operation.",
     "100% UPTIME · 60ms LATENCY · SELF-MONITORING"),

    (18, "Close", "4:58 – 5:10", "01-home.jpg",
     "That's RM ONE. Not a project tool. An operational command centre. "
     "The pipeline, the workforce, the AI, the risk feed, "
     "the costs — every decision LiRo makes about its operation, in one "
     "live connected system.",
     "RM ONE · ONE COMMAND CENTRE · ONE OPERATION"),
]


def style_title(p, text, size=22, color=(255, 255, 255), bg_dark=True, align=None):
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)


def style_body(p, text, size=11, color=(40, 44, 52), bold=False, italic=False):
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)


def style_label(p, label, value, label_color=(86, 156, 86)):
    r1 = p.add_run(label)
    r1.font.name = "Calibri"
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(*label_color)
    r2 = p.add_run("  " + value)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(80, 86, 96)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # === COVER ===
    title = doc.add_paragraph()
    style_title(title, "RM ONE", size=44, color=(86, 156, 86))
    sub = doc.add_paragraph()
    style_title(sub, "Operational Command Centre", size=20, color=(40, 44, 52))
    tag = doc.add_paragraph()
    style_body(tag, "HeyGen Avatar IV Walkthrough Script", size=13,
               color=(100, 108, 120), italic=True)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    style_label(meta, "TENANT", "Liro_Poc")
    meta.add_run("     ")
    style_label(meta, "USER", "Administrator_Liro_Poc")
    meta.add_run("     ")
    style_label(meta, "THEME", "Dark Mode (real logged-in data)")
    meta2 = doc.add_paragraph()
    style_label(meta2, "RUNTIME", "~5 minutes 10 seconds")
    meta2.add_run("     ")
    style_label(meta2, "SCENES", "18")
    meta2.add_run("     ")
    style_label(meta2, "VOICE", "Confident operator · 155–165 wpm")

    doc.add_paragraph()

    # === POSITIONING NOTE ===
    pn = doc.add_paragraph()
    style_title(pn, "Positioning", size=14, color=(86, 156, 86))
    pos = doc.add_paragraph()
    style_body(pos,
        "RM ONE is positioned throughout this script as an operational command "
        "centre — the single live surface where every signal across LiRo's "
        "business arrives and turns into a decision. It is deliberately NOT "
        "described as a project management tool, a scheduling tool, or a "
        "dashboard. Vocabulary across all 18 scenes uses: ",
        size=11)
    style_body(pos,
        "command centre, mission control, operational pulse, operators, "
        "signals, risk feed, ground truth, forward visibility, "
        "self-monitoring infrastructure.",
        size=11, italic=True, color=(86, 156, 86))

    doc.add_page_break()

    # === SCENES ===
    for (num, title_text, runtime, shot, vo, ost) in SCENES:
        # Scene header
        h = doc.add_paragraph()
        style_title(h, f"Scene {num:02d}  ·  {title_text}",
                    size=16, color=(86, 156, 86))

        # Runtime + shot meta
        meta_p = doc.add_paragraph()
        style_label(meta_p, "RUNTIME", runtime)
        meta_p.add_run("        ")
        style_label(meta_p, "B-ROLL", f"screenshots/{shot}")

        # Image
        img_path = SHOTS / shot
        if img_path.exists():
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.add_run().add_picture(str(img_path), width=Inches(6.5))
        else:
            err = doc.add_paragraph()
            style_body(err, f"[missing: {img_path}]", size=10,
                       color=(180, 60, 60), italic=True)

        # Voiceover
        vo_label = doc.add_paragraph()
        style_label(vo_label, "VOICEOVER", "")
        vo_p = doc.add_paragraph()
        style_body(vo_p, f"\u201c{vo}\u201d", size=12,
                   color=(40, 44, 52))

        # On-screen text
        ost_label = doc.add_paragraph()
        style_label(ost_label, "ON-SCREEN TEXT", ost,
                    label_color=(180, 120, 40))

        doc.add_paragraph()
        if num < len(SCENES):
            doc.add_page_break()

    # === PRODUCTION NOTES ===
    doc.add_page_break()
    pn = doc.add_paragraph()
    style_title(pn, "Production Notes for HeyGen",
                size=18, color=(86, 156, 86))
    notes = [
        ("Pace", "155–165 words per minute. Half-second pause between scenes."),
        ("Tone", "Confident operator — not salesy, not technical. Avatar IV's "
                 "natural cadence is the right register."),
        ("Captions",
         "Burn-in the ON-SCREEN TEXT line from each scene as a lower-third "
         "or chyron synced to the scene's first 3 seconds."),
        ("Music",
         "Sparse synth pad, ~60 BPM, ducked under VO. Lift at Scenes 12 "
         "(AI) and 19 (close)."),
        ("Cuts",
         "0.8s crossfade between Resource sub-tab scenes (8→9→10→11). "
         "Hard cut elsewhere."),
        ("Vocabulary discipline",
         "Never use the words \u201cproject management,\u201d \u201ctracker,\u201d "
         "\u201cscheduling tool,\u201d or \u201cdashboard.\u201d Always: "
         "command centre, operational pulse, signals, operators."),
        ("Asset frame",
         "All 18 screenshots are 1280\u00d7720, dark mode, real logged-in "
         "LiRo data — drop in as full-frame B-roll."),
        ("Re-shoot credentials",
         "Tenant Liro_Poc · User Administrator_Liro_Poc · Password "
         "Hilary2023! · Theme defaults to dark."),
    ]
    for k, v in notes:
        p = doc.add_paragraph()
        r = p.add_run(k + ":  ")
        r.font.name = "Calibri"; r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = RGBColor(86, 156, 86)
        r2 = p.add_run(v)
        r2.font.name = "Calibri"; r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(40, 44, 52)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
